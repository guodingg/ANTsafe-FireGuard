"""
报告生成服务 - 支持多种格式
"""

import io
import os
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Optional

from secscan.database import async_session_maker
from secscan.models.scan import ScanTask, TaskStatus
from secscan.models.asset import Asset
from secscan.models.vuln import Vulnerability
from secscan.models.report import Report, ReportType
from secscan.config import settings

class ReportGenerator:
    """报告生成器"""
    
    # 严重性颜色映射
    SEVERITY_COLORS = {
        "critical": (220, 53, 69, 0.1),      # 红色
        "high": (255, 140, 0, 0.1),         # 橙色
        "medium": (255, 193, 7, 0.1),      # 黄色
        "low": (40, 167, 69, 0.1),          # 绿色
        "info": (0, 123, 255, 0.1)          # 蓝色
    }
    
    @classmethod
    async def generate_report(
        cls,
        task_id: int,
        user_id: int,
        report_type: ReportType = ReportType.MARKDOWN
    ) -> Report:
        """生成报告"""
        
        async with async_session_maker() as db:
            from sqlalchemy import select
            
            # 获取任务
            result = await db.execute(select(ScanTask).where(ScanTask.id == task_id))
            task = result.scalar_one_or_none()
            
            if not task:
                raise ValueError(f"任务 {task_id} 不存在")
            
            # 获取资产
            assets_result = await db.execute(
                select(Asset).where(Asset.task_id == task_id)
            )
            assets = assets_result.scalars().all()
            
            # 获取漏洞
            vulns_result = await db.execute(
                select(Vulnerability).where(Vulnerability.task_id == task_id)
            )
            vulns = vulns_result.scalars().all()
            
            # 根据类型生成
            if report_type == ReportType.MARKDOWN:
                content = await cls._generate_markdown(task, assets, vulns)
                file_size = len(content.encode('utf-8'))
            elif report_type == ReportType.HTML:
                content = await cls._generate_html(task, assets, vulns)
                file_size = len(content.encode('utf-8'))
            elif report_type == ReportType.PDF:
                content = await cls._generate_pdf(task, assets, vulns)
                file_size = len(content)
            elif report_type == ReportType.WORD:
                content = await cls._generate_word(task, assets, vulns)
                file_size = len(content)
            elif report_type == ReportType.EXCEL:
                content = await cls._generate_excel(task, assets, vulns)
                file_size = len(content)
            else:
                raise ValueError(f"不支持的报告格式: {report_type}")
            
            # 保存报告
            report = Report(
                task_id=task_id,
                user_id=user_id,
                name=f"{task.name} - 安全评估报告",
                type=report_type,
                content=content.decode('utf-8') if isinstance(content, bytes) else content,
                file_size=file_size
            )
            db.add(report)
            await db.commit()
            await db.refresh(report)
            
            return report
    
    @classmethod
    async def _generate_markdown(cls, task: ScanTask, assets: list, vulns: list) -> str:
        """生成Markdown报告"""
        
        # 统计
        alive_assets = [a for a in assets if a.status == "alive"]
        critical_vulns = [v for v in vulns if v.severity.value == "critical"]
        high_vulns = [v for v in vulns if v.severity.value == "high"]
        medium_vulns = [v for v in vulns if v.severity.value == "medium"]
        low_vulns = [v for v in vulns if v.severity.value == "low"]
        
        lines = [
            f"# {task.name} - 安全评估报告",
            "",
            f"**生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            f"**扫描目标**: {task.target}",
            f"**扫描类型**: {task.scan_type.value}",
            "",
            "## 执行摘要",
            "",
            f"| 指标 | 数值 |",
            f"|------|------|",
            f"| 扫描时间 | {task.started_at.strftime('%Y-%m-%d %H:%M:%S') if task.started_at else 'N/A'} - {task.finished_at.strftime('%Y-%m-%d %H:%M:%S') if task.finished_at else 'N/A'} |",
            f"| 发现资产 | {len(alive_assets)} 个存活 / {len(assets)} 个总资产 |",
            f"| 发现漏洞 | {len(vulns)} 个 |",
            f"| 严重漏洞 | {len(critical_vulns)} 个 |",
            f"| 高危漏洞 | {len(high_vulns)} 个 |",
            f"| 中危漏洞 | {len(medium_vulns)} 个 |",
            f"| 低危漏洞 | {len(low_vulns)} 个 |",
            "",
            "## 风险评估",
            "",
            cls._generate_risk_assessment(len(critical_vulns), len(high_vulns), len(assets)),
            "",
            "## 漏洞详情",
            ""
        ]
        
        # 按严重性分组显示
        for severity, vulns_list in [
            ("严重", critical_vulns),
            ("高危", high_vulns),
            ("中危", medium_vulns),
            ("低危", low_vulns)
        ]:
            if vulns_list:
                lines.append(f"### {severity} - {len(vulns_list)} 个")
                lines.append("")
                for vuln in vulns_list:
                    lines.append(f"#### {vuln.name}")
                    lines.append(f"- **CVE**: {vuln.cve or 'N/A'}")
                    lines.append(f"- **严重性**: {vuln.severity.value}")
                    lines.append(f"- **描述**: {vuln.description or '无'}")
                    if vuln.remediation:
                        lines.append(f"- **修复建议**: {vuln.remediation}")
                    lines.append("")
        
        # 资产清单
        lines.extend([
            "## 资产清单",
            "",
            f"| IP地址 | 端口 | 服务 | 产品 | 版本 | 状态 |",
            f"|--------|------|------|------|------|------|"
        ])
        
        for asset in assets:
            lines.append(f"| {asset.ip} | {asset.port} | {asset.service} | {asset.product} | {asset.version} | {asset.status} |")
        
        lines.extend([
            "",
            "---",
            "",
            f"© {settings.COPYRIGHT}"
        ])
        
        return "\n".join(lines)
    
    @classmethod
    async def _generate_html(cls, task: ScanTask, assets: list, vulns: list) -> str:
        """生成HTML报告"""
        
        # 统计
        severity_counts = {
            "critical": sum(1 for v in vulns if v.severity.value == "critical"),
            "high": sum(1 for v in vulns if v.severity.value == "high"),
            "medium": sum(1 for v in vulns if v.severity.value == "medium"),
            "low": sum(1 for v in vulns if v.severity.value == "low")
        }
        
        html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{task.name} - 安全评估报告</title>
    <style>
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; line-height: 1.6; color: #333; }}
        .container {{ max-width: 1200px; margin: 0 auto; padding: 20px; }}
        .header {{ background: linear-gradient(135deg, #1677FF 0%, #52C41A 100%); color: white; padding: 40px 20px; text-align: center; }}
        .header h1 {{ font-size: 28px; margin-bottom: 10px; }}
        .summary {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 20px; margin: 30px 0; }}
        .stat-card {{ background: white; border-radius: 8px; padding: 20px; box-shadow: 0 2px 8px rgba(0,0,0,0.1); text-align: center; }}
        .stat-card .value {{ font-size: 36px; font-weight: bold; color: #1677FF; }}
        .stat-card .label {{ color: #666; margin-top: 5px; }}
        .section {{ background: white; border-radius: 8px; padding: 20px; margin: 20px 0; box-shadow: 0 2px 8px rgba(0,0,0,0.1); }}
        .section h2 {{ color: #1677FF; border-bottom: 2px solid #1677FF; padding-bottom: 10px; margin-bottom: 20px; }}
        table {{ width: 100%; border-collapse: collapse; }}
        th, td {{ padding: 12px; text-align: left; border-bottom: 1px solid #eee; }}
        th {{ background: #f5f5f5; font-weight: 600; }}
        .severity-critical {{ color: #dc3545; font-weight: bold; }}
        .severity-high {{ color: #fd7e14; font-weight: bold; }}
        .severity-medium {{ color: #ffc107; font-weight: bold; }}
        .severity-low {{ color: #28a745; font-weight: bold; }}
        .footer {{ text-align: center; padding: 20px; color: #666; font-size: 14px; }}
    </style>
</head>
<body>
    <div class="header">
        <h1>{task.name}</h1>
        <p>安全评估报告</p>
        <p>生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
    </div>
    
    <div class="container">
        <div class="summary">
            <div class="stat-card">
                <div class="value">{len(assets)}</div>
                <div class="label">发现资产</div>
            </div>
            <div class="stat-card">
                <div class="value">{len(vulns)}</div>
                <div class="label">发现漏洞</div>
            </div>
            <div class="stat-card">
                <div class="value">{severity_counts['critical']}</div>
                <div class="label">严重漏洞</div>
            </div>
            <div class="stat-card">
                <div class="value">{severity_counts['high']}</div>
                <div class="label">高危漏洞</div>
            </div>
        </div>
        
        <div class="section">
            <h2>漏洞详情</h2>
            <table>
                <tr>
                    <th>漏洞名称</th>
                    <th>CVE</th>
                    <th>严重性</th>
                    <th>描述</th>
                </tr>"""
        
        for vuln in vulns:
            sev_class = f"severity-{vuln.severity.value}"
            html += f"""
                <tr>
                    <td>{vuln.name}</td>
                    <td>{vuln.cve or 'N/A'}</td>
                    <td class="{sev_class}">{vuln.severity.value.upper()}</td>
                    <td>{vuln.description or '无'[:50]}...</td>
                </tr>"""
        
        html += """
            </table>
        </div>
        
        <div class="section">
            <h2>资产清单</h2>
            <table>
                <tr>
                    <th>IP地址</th>
                    <th>端口</th>
                    <th>服务</th>
                    <th>产品</th>
                    <th>版本</th>
                    <th>状态</th>
                </tr>"""
        
        for asset in assets:
            status_color = "#28a745" if asset.status == "alive" else "#dc3545"
            html += f"""
                <tr>
                    <td>{asset.ip}</td>
                    <td>{asset.port}</td>
                    <td>{asset.service}</td>
                    <td>{asset.product}</td>
                    <td>{asset.version}</td>
                    <td style="color: {status_color}">{asset.status}</td>
                </tr>"""
        
        html += f"""
            </table>
        </div>
    </div>
    
    <div class="footer">
        <p>© {settings.COPYRIGHT}</p>
    </div>
</body>
</html>"""
        
        return html
    
    @classmethod
    async def _generate_pdf(cls, task: ScanTask, assets: list, vulns: list) -> bytes:
        """生成PDF报告"""
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        from reportlab.lib.colors import HexColor
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=2*cm, leftMargin=2*cm)
        
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=24, spaceAfter=30, textColor=HexColor('#1677FF'))
        heading_style = ParagraphStyle('Heading', parent=styles['Heading2'], fontSize=14, spaceBefore=20, spaceAfter=10, textColor=HexColor('#1677FF'))
        
        elements = []
        
        # 标题
        elements.append(Paragraph(task.name, title_style))
        elements.append(Paragraph(f"安全评估报告 - 生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
        elements.append(Spacer(1, 20))
        
        # 统计摘要
        elements.append(Paragraph("执行摘要", heading_style))
        
        summary_data = [
            ["指标", "数值"],
            ["发现资产", str(len(assets))],
            ["存活资产", str(sum(1 for a in assets if a.status == 'alive'))],
            ["发现漏洞", str(len(vulns))],
            ["严重漏洞", str(sum(1 for v in vulns if v.severity.value == 'critical'))],
            ["高危漏洞", str(sum(1 for v in vulns if v.severity.value == 'high'))],
        ]
        
        summary_table = Table(summary_data, colWidths=[5*cm, 5*cm])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), HexColor('#1677FF')),
            ('TEXTCOLOR', (0, 0), (-1, 0), HexColor('#ffffff')),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), HexColor('#f5f5f5')),
            ('GRID', (0, 0), (-1, -1), 1, HexColor('#dddddd')),
        ]))
        elements.append(summary_table)
        elements.append(Spacer(1, 20))
        
        # 漏洞详情
        if vulns:
            elements.append(Paragraph("漏洞详情", heading_style))
            
            vuln_data = [["漏洞名称", "CVE", "严重性"]]
            for vuln in vulns[:50]:  # PDF限制显示50条
                vuln_data.append([
                    vuln.name[:30],
                    vuln.cve or 'N/A',
                    vuln.severity.value.upper()
                ])
            
            vuln_table = Table(vuln_data, colWidths=[9*cm, 3*cm, 3*cm])
            vuln_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), HexColor('#FF4D4F')),
                ('TEXTCOLOR', (0, 0), (-1, 0), HexColor('#ffffff')),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                ('GRID', (0, 0), (-1, -1), 1, HexColor('#dddddd')),
            ]))
            elements.append(vuln_table)
        
        # 资产清单
        if assets:
            elements.append(PageBreak())
            elements.append(Paragraph("资产清单", heading_style))
            
            asset_data = [["IP地址", "端口", "服务", "状态"]]
            for asset in assets[:50]:  # PDF限制显示50条
                asset_data.append([
                    asset.ip,
                    str(asset.port),
                    asset.service,
                    asset.status
                ])
            
            asset_table = Table(asset_data, colWidths=[4*cm, 2*cm, 4*cm, 2*cm])
            asset_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), HexColor('#52C41A')),
                ('TEXTCOLOR', (0, 0), (-1, 0), HexColor('#ffffff')),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                ('GRID', (0, 0), (-1, -1), 1, HexColor('#dddddd')),
            ]))
            elements.append(asset_table)
        
        # 页脚
        elements.append(Spacer(1, 30))
        elements.append(Paragraph(f"© {settings.COPYRIGHT}", styles['Normal']))
        
        doc.build(elements)
        return buffer.getvalue()
    
    @classmethod
    async def _generate_word(cls, task: ScanTask, assets: list, vulns: list) -> bytes:
        """生成Word报告"""
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        
        doc = Document()
        
        # 标题
        title = doc.add_heading(task.name, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph(f"安全评估报告")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        time_para = doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # 执行摘要
        doc.add_heading("执行摘要", level=1)
        
        summary_table = doc.add_table(rows=6, cols=2)
        summary_table.style = 'Table Grid'
        
        summary_data = [
            ("发现资产", str(len(assets))),
            ("存活资产", str(sum(1 for a in assets if a.status == 'alive'))),
            ("发现漏洞", str(len(vulns))),
            ("严重漏洞", str(sum(1 for v in vulns if v.severity.value == 'critical'))),
            ("高危漏洞", str(sum(1 for v in vulns if v.severity.value == 'high'))),
            ("中危漏洞", str(sum(1 for v in vulns if v.severity.value == 'medium'))),
        ]
        
        for i, (label, value) in enumerate(summary_data):
            summary_table.rows[i].cells[0].text = label
            summary_table.rows[i].cells[1].text = value
        
        doc.add_paragraph()
        
        # 漏洞详情
        if vulns:
            doc.add_heading("漏洞详情", level=1)
            
            vuln_table = doc.add_table(rows=len(vulns[:100])+1, cols=4)
            vuln_table.style = 'Table Grid'
            
            # 表头
            header_cells = vuln_table.rows[0].cells
            header_cells[0].text = "漏洞名称"
            header_cells[1].text = "CVE"
            header_cells[2].text = "严重性"
            header_cells[3].text = "描述"
            
            for i, vuln in enumerate(vulns[:100]):
                cells = vuln_table.rows[i+1].cells
                cells[0].text = vuln.name
                cells[1].text = vuln.cve or 'N/A'
                cells[2].text = vuln.severity.value.upper()
                cells[3].text = (vuln.description or '无')[:50]
        
        doc.add_paragraph()
        
        # 资产清单
        if assets:
            doc.add_heading("资产清单", level=1)
            
            asset_table = doc.add_table(rows=len(assets[:100])+1, cols=5)
            asset_table.style = 'Table Grid'
            
            # 表头
            header_cells = asset_table.rows[0].cells
            header_cells[0].text = "IP地址"
            header_cells[1].text = "端口"
            header_cells[2].text = "服务"
            header_cells[3].text = "产品"
            header_cells[4].text = "状态"
            
            for i, asset in enumerate(assets[:100]):
                cells = asset_table.rows[i+1].cells
                cells[0].text = asset.ip
                cells[1].text = str(asset.port)
                cells[2].text = asset.service
                cells[3].text = asset.product
                cells[4].text = asset.status
        
        # 页脚
        doc.add_paragraph()
        footer = doc.add_paragraph(f"© {settings.COPYRIGHT}")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 保存到BytesIO
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
    
    @classmethod
    async def _generate_excel(cls, task: ScanTask, assets: list, vulns: list) -> bytes:
        """生成Excel报告"""
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        
        wb = openpyxl.Workbook()
        
        # 样式定义
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="1677FF", end_color="1677FF", fill_type="solid")
        critical_fill = PatternFill(start_color="FF4D4F", end_color="FF4D4F", fill_type="solid")
        high_fill = PatternFill(start_color="FF8C00", end_color="FF8C00", fill_type="solid")
        medium_fill = PatternFill(start_color="FFC107", end_color="FFC107", fill_type="solid")
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        # Sheet 1: 摘要
        ws_summary = wb.active
        ws_summary.title = "执行摘要"
        
        ws_summary['A1'] = task.name
        ws_summary['A2'] = f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        
        ws_summary['A4'] = "指标"
        ws_summary['B4'] = "数值"
        ws_summary['A4'].font = header_font
        ws_summary['B4'].font = header_font
        ws_summary['A4'].fill = header_fill
        ws_summary['B4'].fill = header_fill
        
        summary_data = [
            ("发现资产", len(assets)),
            ("存活资产", sum(1 for a in assets if a.status == 'alive')),
            ("发现漏洞", len(vulns)),
            ("严重漏洞", sum(1 for v in vulns if v.severity.value == 'critical')),
            ("高危漏洞", sum(1 for v in vulns if v.severity.value == 'high')),
            ("中危漏洞", sum(1 for v in vulns if v.severity.value == 'medium')),
            ("低危漏洞", sum(1 for v in vulns if v.severity.value == 'low')),
        ]
        
        for i, (label, value) in enumerate(summary_data, start=5):
            ws_summary[f'A{i}'] = label
            ws_summary[f'B{i}'] = value
        
        # Sheet 2: 漏洞列表
        ws_vulns = wb.create_sheet("漏洞列表")
        
        vuln_headers = ["漏洞名称", "CVE", "严重性", "描述", "修复建议"]
        for col, header in enumerate(vuln_headers, start=1):
            cell = ws_vulns.cell(row=1, column=col)
            cell.value = header
            cell.font = header_font
            cell.fill = header_fill
        
        for row, vuln in enumerate(vulns, start=2):
            ws_vulns.cell(row=row, column=1).value = vuln.name
            ws_vulns.cell(row=row, column=2).value = vuln.cve or 'N/A'
            ws_vulns.cell(row=row, column=3).value = vuln.severity.value.upper()
            ws_vulns.cell(row=row, column=4).value = vuln.description or '无'
            ws_vulns.cell(row=row, column=5).value = vuln.remediation or '无'
            
            # 根据严重性设置颜色
            sev_cell = ws_vulns.cell(row=row, column=3)
            if vuln.severity.value == 'critical':
                sev_cell.fill = critical_fill
                sev_cell.font = Font(bold=True, color="FFFFFF")
            elif vuln.severity.value == 'high':
                sev_cell.fill = high_fill
                sev_cell.font = Font(bold=True, color="FFFFFF")
            elif vuln.severity.value == 'medium':
                sev_cell.fill = medium_fill
        
        # Sheet 3: 资产列表
        ws_assets = wb.create_sheet("资产列表")
        
        asset_headers = ["IP地址", "端口", "服务", "产品", "版本", "状态"]
        for col, header in enumerate(asset_headers, start=1):
            cell = ws_assets.cell(row=1, column=col)
            cell.value = header
            cell.font = header_font
            cell.fill = header_fill
        
        for row, asset in enumerate(assets, start=2):
            ws_assets.cell(row=row, column=1).value = asset.ip
            ws_assets.cell(row=row, column=2).value = asset.port
            ws_assets.cell(row=row, column=3).value = asset.service
            ws_assets.cell(row=row, column=4).value = asset.product
            ws_assets.cell(row=row, column=5).value = asset.version
            ws_assets.cell(row=row, column=6).value = asset.status
        
        # 调整列宽
        for ws in [ws_summary, ws_vulns, ws_assets]:
            for col in ws.columns:
                max_length = 0
                column = col[0].column_letter
                for cell in col:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                ws.column_dimensions[column].width = adjusted_width
        
        # 保存
        buffer = io.BytesIO()
        wb.save(buffer)
        return buffer.getvalue()
    
    @classmethod
    def _generate_risk_assessment(cls, critical: int, high: int, total_assets: int) -> str:
        """生成风险评估"""
        if critical > 0:
            return f"🔴 **极高风险** - 发现 {critical} 个严重漏洞，建议立即修复"
        elif high > 5:
            return f"🟠 **高风险** - 发现 {high} 个高危漏洞，建议尽快修复"
        elif high > 0:
            return f"🟡 **中风险** - 发现 {high} 个高危漏洞，建议安排修复"
        elif total_assets > 10:
            return f"🟢 **低风险** - 暂未发现高危漏洞"
        else:
            return f"✅ **安全** - 未发现明显安全问题"
